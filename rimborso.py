import fitz 
import pathlib
import os
import re
import openpyxl
from openpyxl.styles import Font, Alignment
import docx


# Parametri iniziali
nome_collaboratore = "Elisa Bianchi"
cartella_viaggi = pathlib.Path.cwd() / "Viaggi da rimborsare"
file_riepilogo = pathlib.Path.cwd() / "Riepilogo rimborsi.xlsx"

# estrazione dati
pattern_partenza_data = r"Stazione di Partenza\s+(.*?)\s*\n+Ore\s.*?-\s*(\d{2}/\d{2}/\d{4})"
pattern_arrivo = r"Stazione di Arrivo\s*(.+?)\nOre"
pattern_importo = r"Importo.*:\s*(\d+\.\d{2})"

# Creazione/apertura file excel
if file_riepilogo.exists():
    riepilogo_xlsx = openpyxl.load_workbook(file_riepilogo)
    nuovo_foglio = riepilogo_xlsx.create_sheet()
else:
    riepilogo_xlsx = openpyxl.Workbook()
    nuovo_foglio = riepilogo_xlsx.active

# Intestazioni
intestazioni = ["Data viaggio", "Stazione di partenza", "Stazione di arrivo", "Prezzo biglietto", "Nome file"]
for col, titolo in zip("ABCDE", intestazioni):
    nuovo_foglio[f"{col}1"] = titolo
    nuovo_foglio[f"{col}1"].font = Font(bold=True)

# Ricevuta file Word
ricevuta = docx.Document()
ricevuta.add_heading("Ricevuta rimborsi", 1).alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

# PDF
riga_excel = 2
totale = 0.0

for nome_file in os.listdir(cartella_viaggi):
    if not nome_file.lower().endswith(".pdf"):
        continue

    percorso_file = cartella_viaggi / nome_file
    doc = fitz.open(percorso_file)

    for pagina in doc:
        
        partenza = arrivo = data = prezzo = None

        testo = pagina.get_text()

        match_partenza_data = re.search(pattern_partenza_data, testo)
        match_arrivo = re.search(pattern_arrivo, testo)
        match_importo = re.search(pattern_importo, testo)

        if match_partenza_data:
            partenza = match_partenza_data.group(1).strip()
            data = match_partenza_data.group(2).strip()
        if match_arrivo:
            arrivo = match_arrivo.group(1).strip()
        if match_importo:
            prezzo = float(match_importo.group(1).replace(",", "."))

        if partenza and arrivo and data and prezzo is not None:
            nuovo_foglio[f"A{riga_excel}"] = data
            nuovo_foglio[f"B{riga_excel}"] = partenza
            nuovo_foglio[f"C{riga_excel}"] = arrivo
            nuovo_foglio[f"D{riga_excel}"] = prezzo
            nuovo_foglio[f"D{riga_excel}"].number_format = '€ 0.00'
            nuovo_foglio[f"E{riga_excel}"] = nome_file
            totale += prezzo
            riga_excel += 1

# Inserimento riga totale
nuovo_foglio[f"C{riga_excel}"] = "Totale:"
nuovo_foglio[f"C{riga_excel}"].alignment = Alignment(horizontal="right")
nuovo_foglio[f"C{riga_excel}"].font = Font(bold=True)
nuovo_foglio[f"D{riga_excel}"] = f"=SUM(D2:D{riga_excel-1})"
nuovo_foglio[f"D{riga_excel}"].font = Font(bold=True)
nuovo_foglio[f"D{riga_excel}"].number_format = '€ 0.00'

# Salvataggio Excel
riepilogo_xlsx.save(file_riepilogo)

# ricevuta word
ricevuta.add_paragraph(
    f"Il/la sottoscritto/a {nome_collaboratore} dichiara di ricevere la somma di € {totale:.2f} a titolo di rimborso spese"
)

for biglietto in nuovo_foglio.iter_rows(min_row=2, max_row=riga_excel-1, values_only=True):
    data, partenza, arrivo, prezzo, _ = biglietto
    ricevuta.add_paragraph(f"{data}: {partenza} - {arrivo} € {prezzo:.2f}", style="List Bullet")

nome_file_ricevuta = f"Ricevuta_rimborso_{nome_collaboratore.replace(' ', '_')}.docx"
ricevuta.save(nome_file_ricevuta)

print(f"Riepilogo Excel salvato in: {file_riepilogo}")
print(f"Ricevuta Word salvata in: {nome_file_ricevuta}")

